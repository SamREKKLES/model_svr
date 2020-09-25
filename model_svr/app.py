import base64

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from nibabel.viewers import OrthoSlicer3D
from scipy import misc

from utils import common

import os
from datetime import datetime
from io import BytesIO
import matplotlib

matplotlib.use('Agg')
import matplotlib.pyplot as plt
import nibabel as nib
import numpy as np
from flask import Flask, request, send_from_directory, session, send_file, make_response, Response
from flask_cors import CORS, cross_origin
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import null
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

from flask_wtf import CSRFProtect
from flask_socketio import SocketIO
import uuid
from sklearn import metrics

from utils.auths import login_required
from utils.common import SQLALCHEMY_DATABASE_URI, failReturn, successReturn, emailSent
from stage1_2 import stage1_init, stage2_init, stage2, load_imgs, stage1_2, to_nii
from flasgger import Swagger

app = Flask(__name__)
APP_ROOT = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(APP_ROOT, 'uploads')
RESULT_FOLDER = os.path.join(APP_ROOT, 'results')
DOC_FOLDER = os.path.join(APP_ROOT, 'doc')
GET_PIC = os.path.join(APP_ROOT, 'pic')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(RESULT_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['RESULT_FOLDER'] = RESULT_FOLDER
app.config['DOC_FOLDER'] = DOC_FOLDER
app.config['GET_PIC'] = GET_PIC
app.config['SQLALCHEMY_DATABASE_URI'] = SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = True
app.config['WTF_CSRF_ENABLED'] = False
db = SQLAlchemy(app)
app.secret_key = os.urandom(24)

# swagger接口文档
swagger_config = Swagger.DEFAULT_CONFIG
swagger_config['title'] = common.SWAGGER_TITLE
swagger_config['description'] = common.SWAGGER_DESC
Swagger(app, config=swagger_config)

CSRFProtect(app)
# enable CORS
CORS(app, supports_credentials=True, resources={r'/*': {'origins': '*'}})
socketio = SocketIO(app, cors_allowed_origins='*')


class Img(db.Model):
    __tablename__ = 'imgs'
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), unique=True)
    uploadname = db.Column(db.String(255), unique=False)
    timestamp = db.Column(db.DateTime)
    type = db.Column(db.String(255), unique=False)
    patient_id = db.Column(db.Integer)
    doctor_id = db.Column(db.Integer)

    def __init__(self, filename, uploadname, img_type, patient, doctor):
        self.filename = filename
        self.type = img_type
        self.patient_id = patient
        self.uploadname = uploadname
        self.doctor_id = doctor
        self.timestamp = datetime.now()

    def __repr__(self):
        return '<DWI %r>' % self.filename


class Result(db.Model):
    __tablename__ = 'results'
    id = db.Column(db.Integer, primary_key=True)
    filename1 = db.Column(db.String(255), unique=True)
    filename2 = db.Column(db.String(255), unique=True)
    timestamp = db.Column(db.DateTime)
    modeltype = db.Column(db.String(255), unique=False)
    dwi_name = db.Column(db.String(255), unique=False)
    adc_name = db.Column(db.String(255), unique=False)
    info = db.Column(db.Float, unique=False)
    patient_id = db.Column(db.Integer)
    doctor_id = db.Column(db.Integer)
    realimg = db.Column(db.String(255), unique=True)
    roi = db.Column(db.String(255), unique=True)
    size = db.Column(db.Float, unique=False)

    def __init__(self, filename1, filename2, modeltype, patient, doctor, dwi_name, adc_name, info, size):
        self.filename1 = filename1
        self.filename2 = filename2
        self.modeltype = modeltype
        self.patient_id = patient
        self.doctor_id = doctor
        self.dwi_name = dwi_name
        self.adc_name = adc_name
        self.info = info
        self.timestamp = datetime.now()
        self.size = size

    def __repr__(self):
        return '<DWI %r>' % self.filename


class Patient(db.Model):
    __tablename__ = 'patients'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(255), unique=True)
    age = db.Column(db.Integer)
    sex = db.Column(db.Integer)
    record_id = db.Column(db.String(255))
    info = db.Column(db.String(255))
    result = db.Column(db.String(255))
    cva = db.Column(db.String(255))
    state = db.Column(db.String(255))
    create_time = db.Column(db.DateTime)
    update_time = db.Column(db.DateTime)
    doctor_id = db.Column(db.Integer)

    def __init__(self, username, recordID, state, doctor, age, sex, info, result, cva):
        self.username = username
        self.doctor_id = doctor
        self.record_id = recordID
        self.state = state
        self.age = age
        self.sex = sex
        self.info = info
        self.result = result
        self.cva = cva
        self.create_time = datetime.now()
        self.update_time = self.create_time

    def __repr__(self):
        return '<Patient %r>' % self.username


class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(255), unique=True)
    password = db.Column(db.String(255))
    realname = db.Column(db.String(255), unique=False)
    userType = db.Column(db.Integer)

    def __init__(self, username, password, realname, userType=3):
        password = generate_password_hash(password)
        self.username = username
        self.password = password
        self.realname = realname
        self.userType = userType

    def check_password(self, password):
        return check_password_hash(self.password, password)

    def to_json(self):
        return {'id': self.id, 'username': self.username,
                'realname': self.realname, 'usertype': self.userType}

    def get_id(self):
        return str(self.id)

    def __repr__(self):
        return '<User %r>' % self.username


db.create_all()


def _get_current_user():
    """
    获取当前用户
      User
    """
    currentID = session["user_id"]
    if currentID:
        return User.query.filter_by(id=currentID).first()
    return null


def add_item(id, img_type, filename, uploadname):
    """
    增加img信息
    :param id:
    :param img_type:
    :param filename:
    :param uploadname:
      boolean
    """
    patient = Patient.query.filter_by(id=id).first()
    doctorID = session["user_id"]
    if patient is None:
        return False
    ct = Img(filename, uploadname, img_type, patient.id, doctorID)
    db.session.add(ct)
    db.session.commit()
    return True


def img_to_base64(img):
    """
    img转换
    :param img:
      string
    """
    output_buffer = BytesIO()
    plt.imsave(output_buffer, img, cmap='gray')
    byte_data = output_buffer.getvalue()
    base64_data = base64.b64encode(byte_data)
    return "data:image/jpg;base64," + base64_data.decode('ascii')


@app.route('/api/imgUpload', methods=['POST'])
@login_required
@cross_origin()
def img_upload():
    """
      img图像上传，返回base64
    ---
    tags:
      - model_svr API
    parameters:
      - name: file
        in: formData
        type: file
        required: true
        description: The language name
      - name: patientID
        in: formData
        type: integer
        required: true
      - name: type
        in: formData
        type: string
        required: true
        description: ADC或DWI
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      success:
        schema:
          type: object
          properties:
            filename:
                type: string
                example: filename
            imgs:
                type: array
                items:
                    type: string
                    example: [data:image/jpg;base64 xxxxxx, ···]
            slices:
                type: string
                example: 21
        description: 成功
      fail:
        schema:
          type: object
          properties:
            status:
              type: string
              example: fail.
            msg:
              type: string
              example: imgUpload出错 or 图像上传失败,无该用户
            data:
              type: string
              example: error
        description: 失败
    """
    try:
        response_object = {}
        file = request.files['file']
        uploadname = secure_filename(file.filename)
        id = request.form['patientID']
        img_type = request.form['type']
        filename = img_type + "_" + uuid.uuid4().hex + ".nii.gz"
        save_path = app.config['UPLOAD_FOLDER']
        os.makedirs(save_path, exist_ok=True)
        save_file = os.path.join(save_path, filename)
        if not add_item(id, img_type, filename, uploadname):
            return failReturn("", "imgUpload: 图像上传失败,无该用户")
        else:
            file.save(save_file)
        response_object['filename'] = filename
        response_object['imgs'], response_object['slices'] = get_all_slice(save_file)
        return successReturn(response_object, "img: 获取成功")
    except Exception as e:
        return failReturn(format(e), "imgUpload出错")


# def _get_filename(id):
#     """
#     获取img信息
#     :param id:
#       img
#     """
#
#     def to_dict(id):
#         imgs = Img.query.filter_by(patient_id=id).order_by(Img.timestamp).all()[::-1]
#         res = []
#         for r in imgs:
#             res.append({'id': r.id, 'time': r.timestamp, 'filename': r.filename,
#                         'uploadName': r.uploadname, 'type': r.type})
#         return res
#
#     patient = Patient.query.filter_by(id=id).first()
#     doctor = _get_current_user()
#     if doctor.userType == 1 or patient.docter_id == doctor.id:
#         return to_dict(id)
#     else:
#         return None
#
#
# @app.route('/api/getFilename', methods=['POST'])
# @login_required
# @cross_origin()
# def get_filename():
#     """
#     根据病人id获取img信息
#       json
#     ---
#     tags:
#       - model_svr API
#     parameters:
#       - name: body
#         in: body
#         required: true
#         schema:
#           id: 获取img
#           required:
#             - patientID
#           properties:
#             patientID:
#               type: integer
#               description: patientID
#       - name: Authorization
#         in: header
#         type: string
#         required: true
#         description: token
#     responses:
#       success:
#         schema:
#           type: object
#           properties:
#             status:
#               type: string
#               example: success.
#             msg:
#               type: string
#               example: 获取result成功
#             data:
#               type: object
#               properties:
#                   imgs:
#                     type: array
#                     items:
#                         type: object
#                         properties:
#                             id:
#                                 type: string
#                                 example: 1
#                             time:
#                                 type: string
#                                 example: date-time
#                             filename:
#                                 type: string
#                                 example: filename
#                             uploadName:
#                                 type: string
#                                 example: uploadName
#                             type:
#                                 type: string
#                                 example: type
#         description: 成功
#       fail:
#         schema:
#           type: object
#           properties:
#             status:
#               type: string
#               example: fail.
#             msg:
#               type: string
#               example: getfilename出错 or 权限不足或无该result，获取filename失败
#             data:
#               type: string
#               example: error
#         description: 失败
#     """
#     try:
#         json = request.get_json()
#         id = json['patientID']
#         imgs = _get_filename(id)
#         if imgs:
#             return successReturn({"imgs": imgs}, "getResults: 获取result成功")
#         else:
#             return failReturn("", "getResults: 权限不足或无该result，获取filename失败")
#     except Exception as e:
#         return failReturn(format(e), "getfilename出错")


def _get_results(id):
    """
    获取result信息
    :param id:
      result
    """

    def to_dicts(p):
        results = Result.query.filter_by(patient_id=p.id).order_by(Result.timestamp.desc()).all()[::-1]
        res = []
        dwi_file = {}
        adc_file = {}
        perfusion = {}
        non_perfusion = {}
        for r in results:
            dwi_file['dwi_file'] = r.dwi_name
            dwi_file['dwi_imgs'], dwi_file['dwi_slices'] = get_all_slice(r.dwi_name)
            adc_file['adc_file'] = r.adc_name
            adc_file['adc_imgs'], adc_file['adc_slices'] = get_all_slice(r.adc_name)
            perfusion['perfusion_file'] = r.filename1
            perfusion['perfusion_imgs'], perfusion['perfusion_slices'] = get_all_slice(r.filename1)
            non_perfusion['non_perfusion_file'] = r.filename2
            non_perfusion['non_perfusion_imgs'], non_perfusion['non_perfusion_slices'] = get_all_slice(r.filename2)
            res.append({'id': r.id, 'time': r.timestamp, 'dwi': dwi_file, 'adc': adc_file,
                        'perfusion': perfusion, 'non_perfusion': non_perfusion, 'modelType': r.modeltype,
                        "info": r.info, "size_percent": r.size})
        return res

    patient = Patient.query.filter_by(id=id).first()
    doctor = _get_current_user()
    if doctor.userType == 1 or patient.docter_id == doctor.id:
        return to_dicts(patient)
    else:
        return None


@app.route('/api/getResultsByPatient', methods=['POST'])
@login_required
@cross_origin()
def get_results_by_patientID():
    """
      根据病人id获取result，对应影响管理部分的api
    ---
    tags:
      - model_svr API
    parameters:
      - name: body
        in: body
        required: true
        schema:
          id: 获取result
          required:
            - patientID
          properties:
            patientID:
              type: integer
              description: patientID
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      success:
        schema:
          type: object
          properties:
            status:
              type: string
              example: success.
            msg:
              type: string
              example: 获取result成功
            data:
              type: object
              properties:
                  results:
                    type: array
                    items:
                        type: object
                        properties:
                            id:
                                type: string
                                example: 1
                            time:
                                type: string
                                format: date-time
                            dwi:
                                type: object
                                properties:
                                    dwi_file:
                                        type: string
                                        example: dwi_file
                                    dwi_imgs:
                                        type: array
                                        items:
                                            type: string
                                            example: [data:image/jpg;base64 xxxxxx, ···]
                                    dwi_slices:
                                        type: string
                                        example: 21
                            adc:
                                type: object
                                properties:
                                    adc_file:
                                        type: string
                                        example: adc_file
                                    adc_imgs:
                                        type: array
                                        items:
                                            type: string
                                            example: [data:image/jpg;base64 xxxxxx, ···]
                                    adc_slices:
                                        type: string
                                        example: 21
                            perfusion:
                                type: object
                                properties:
                                    perfusion_file:
                                        type: string
                                        example: perfusion_file
                                    perfusion_imgs:
                                        type: array
                                        items:
                                            type: string
                                            example: [data:image/jpg;base64 xxxxxx, ···]
                                    perfusion_slices:
                                        type: string
                                        example: 21
                            non_perfusion:
                                type: object
                                properties:
                                    non_perfusion_file:
                                        type: string
                                        example: non_perfusion_file
                                    non_perfusion_imgs:
                                        type: array
                                        items:
                                            type: string
                                            example: [data:image/jpg;base64 xxxxxx, ···]
                                    non_perfusion_slices:
                                        type: string
                                        example: 21
                            modelType:
                                type: string
                                example: Random Forest
                            info:
                                type: number
                                format: float
                                example: 120.0
                            size_percent:
                                type: number
                                format: float
                                example: 0.5
        description: 成功
      fail:
        schema:
          type: object
          properties:
            status:
              type: string
              example: fail.
            msg:
              type: string
              example: getResults出错 or 权限不足或无该result，获取result失败
            data:
              type: string
              example: error
        description: 失败
    """
    try:
        json = request.get_json()
        id = json['patientID']
        results = _get_results(id)

        if results:
            return successReturn({"results": results}, "getResults: 获取result成功")
        else:
            return failReturn("", "getResults: 权限不足或无该result，获取result失败")
    except Exception as e:
        return failReturn(format(e), "getResults出错")


# def _get_inp_out(id):
#     """
#     获取图像结果信息
#     :param id:
#       adc_file, dwi_file, res_file1, res_file2, info
#     """
#     result = Result.query.filter_by(id=id).first()
#     if not result:
#         return None, None, None, None, None, None, None
#     adc_file = result.adc_name
#     dwi_file = result.dwi_name
#     res_file1 = result.filename1
#     res_file2 = result.filename2
#     info = result.info
#     modelType = result.modeltype
#     timeStamp = result.timestamp
#     return adc_file, dwi_file, res_file1, res_file2, info, modelType, timeStamp
#
#
# @app.route('/api/getInpOut', methods=['POST'])
# @login_required
# @cross_origin()
# def get_inp_out():
#     """
#     获取图像信息
#       json
#     ---
#     tags:
#       - model_svr API
#     parameters:
#       - name: body
#         in: body
#         required: true
#         schema:
#           id: 获取图像信息
#           required:
#             - resultID
#           properties:
#             resultID:
#               type: integer
#               description: resultID
#       - name: Authorization
#         in: header
#         type: string
#         required: true
#         description: token
#     responses:
#       success:
#         schema:
#           type: object
#           properties:
#             status:
#               type: string
#               example: success.
#             msg:
#               type: string
#               example: 成功获取数据信息
#             data:
#               type: object
#               properties:
#                 dwi_file:
#                     type: string
#                     example: filename
#                 dwi_imgs:
#                     type: string
#                     example: base64
#                 dwi_slices:
#                     type: string
#                     example: 21
#                 adc_file:
#                     type: string
#                     example: filename
#                 adc_imgs:
#                     type: string
#                     example: base64
#                 adc_slices:
#                     type: string
#                     example: 21
#                 res_file1:
#                     type: string
#                     example: filename
#                 res_imgs1:
#                     type: string
#                     example: base64
#                 res_slices1:
#                     type: string
#                     example: 21
#                 res_file2:
#                     type: string
#                     example: filename
#                 res_imgs2:
#                     type: string
#                     example: base64
#                 res_slices2:
#                     type: string
#                     example: 21
#                 info:
#                     type: number
#                     format: float
#                     example: 120.0
#         description: 成功
#       fail:
#         schema:
#           type: object
#           properties:
#             status:
#               type: string
#               example: fail.
#             msg:
#               type: string
#               example: 无数据信息 or getInpOut出错
#             data:
#               type: string
#               example: error
#         description: 失败
#     """
#     try:
#         response_object = {}
#         json = request.get_json()
#         id = json['resultID']
#         adc_file, dwi_file, res_file1, res_file2, info, modelType, timeStamp = _get_inp_out(id)
#         if adc_file or dwi_file or res_file1 or res_file2:
#             if dwi_file:
#                 response_object['dwi_file'] = dwi_file
#                 dwi_file = os.path.join(app.config['UPLOAD_FOLDER'], dwi_file)
#                 response_object['dwi_imgs'], response_object['dwi_slices'] = get_all_slice(dwi_file)
#             if adc_file:
#                 response_object['adc_file'] = adc_file
#                 adc_file = os.path.join(app.config['UPLOAD_FOLDER'], adc_file)
#                 response_object['adc_imgs'], response_object['adc_slices'] = get_all_slice(adc_file)
#             if res_file1:
#                 response_object['res_file1'] = res_file1
#                 res_file1 = os.path.join(app.config['RESULT_FOLDER'], res_file1)
#                 response_object['res_imgs1'], response_object['res_slices1'] = get_all_slice(res_file1, thres=0.25)
#             if res_file2:
#                 response_object['res_file2'] = res_file2
#                 res_file2 = os.path.join(app.config['RESULT_FOLDER'], res_file2)
#                 response_object['res_imgs2'], response_object['res_slices2'] = get_all_slice(res_file2, thres=0.25)
#             if info:
#                 info = round(info, 2)
#                 response_object['info'] = info
#             if modelType:
#                 response_object['modelType'] = modelType
#             if timeStamp:
#                 response_object['timeStamp'] = timeStamp
#         else:
#             return failReturn("", "getInpOut: 无数据信息")
#         return successReturn(response_object, "getInpOut: 成功获取数据信息")
#     except Exception as e:
#         return failReturn(format(e), "getInpOut出错")


def get_all_slice(filename, thres=None):
    """
    获取所有切片
    :param filename:
    :param thres:
      res
    """
    if not filename:
        return None, None
    imgs = nib.load(filename).get_fdata()
    imgs = np.squeeze(imgs)
    if thres:
        imgs[imgs >= thres] = 1
        imgs[imgs < thres] = 0
    res = []
    for idx in range(imgs.shape[2]):
        res.append(img_to_base64(imgs[:, :, idx]))
    return res, str(imgs.shape[2])


def _perfImgs(id):
    """
    获取perfImgs信息
    :param id:
      result
    """

    def perfImgs_dicts(p):
        results = Result.query.filter_by(patient_id=p.id).order_by(Result.timestamp.desc()).all()[::-1]
        res = []
        perfusion = {}
        for r in results:
            perfusion['perfusion_file'] = r.filename1
            perfusion['perfusion_imgs'], perfusion['perfusion_slices'] = get_all_slice(r.filename1)
            text_info = str(r.timestamp) + " 梗死区域面积为: " + str(r.size) + " 溶栓治疗获益为: " + str(r.info)
            res.append({'id': r.id, 'perfusion': perfusion, "text_info": text_info})
        return res

    patient = Patient.query.filter_by(id=id).first()
    doctor = _get_current_user()
    if doctor.userType == 1 or patient.docter_id == doctor.id:
        return perfImgs_dicts(patient)
    else:
        return None


@app.route('/api/getPerfImgs', methods=['POST'])
@login_required
@cross_origin()
def get_perfImgs():
    """
      对应图像展窗部分的显示
    ---
    tags:
      - model_svr API
    parameters:
      - name: body
        in: body
        required: true
        schema:
          id: 获取ct结果 dwi或者adc
          properties:
            dwi:
              type: string
              description: dwi
            adc:
              type: string
              description: adc
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      success:
        schema:
          type: object
          properties:
            status:
              type: string
              example: success.
            msg:
              type: string
              example: 获取成功
            data:
              type: object
              properties:
                perf_imgs:
                  type: array
                  items:
                    type: object
                    properties:
                        id:
                            type: string
                            example: 1
                        perfusion:
                            type: object
                            properties:
                                perfusion_file:
                                    type: string
                                    example: dwi_file
                                perfusion_imgs:
                                    type: array
                                    items:
                                        type: string
                                        example: [data:image/jpg;base64 xxxxxx, ···]
                                perfusion_slices:
                                    type: string
                                    example: 21
                        text_info:
                            type: string
                            example: "2020-9-25 00:00:00 梗死区域面积比例为: 0.5 溶栓治疗获益为: 120.0"
        description: 成功
      fail:
        schema:
          type: object
          properties:
            status:
              type: string
              example: fail.
            msg:
              type: string
              example: img出错 or 参数为空
            data:
              type: string
              example: error
        description: 失败
    """
    try:
        response_object = {}
        json = request.get_json()
        id = json['patientID']
        perfImgs = _perfImgs(id)
        if perfImgs:
            response_object['perf_imgs'] = perfImgs
            return successReturn(response_object, "img: 获取成功")
        else:
            return failReturn("", "病人图像结果未找到")
    except Exception as e:
        return failReturn(format(e), "img出错")


@app.route('/api/analyze', methods=['POST'])
@login_required
@cross_origin()
def analyze():
    """
      结果预测"Random Forest" or "Random Forest+U-Net"
    ---
    tags:
      - model_svr API
    parameters:
      - name: body
        in: body
        required: true
        schema:
          id: 结果分析
          required:
            - patientID
            - backmodel
            - dwi_file
            - adc_file
          properties:
            patientID:
              type: integer
              description: patientID
            backmodel:
              type: string
              description: backmodel
            dwi_file:
              type: string
              description: dwi_file
            adc_file:
              type: string
              description: adc_file
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      success:
        schema:
          type: object
          properties:
            status:
              type: string
              example: success.
            msg:
              type: string
              example: 分析成功
            data:
              type: object
              properties:
                perf_res_imgs:
                    type: array
                    items:
                        type: string
                        example: [data:image/jpg;base64 xxxxxx, ···]
                perf_res_slices:
                    type: string
                    example: 21
                nonperf_res_imgs:
                    type: array
                    items:
                        type: string
                        example: [data:image/jpg;base64 xxxxxx, ···]
                nonperf_res_slices:
                    type: string
                    example: 21
                res_path1:
                    type: string
                    example: res_path1
                res_path2:
                    type: string
                    example: res_path2
                info:
                    type: number
                    format: float
                    example: 120.0
                size:
                    type: number
                    format: float
                    example: 0.05
        description: 成功
      fail:
        schema:
          type: object
          properties:
            status:
              type: string
              example: fail.
            msg:
              type: string
              example: 输入参数缺失 or analyze出错, 请选择正确算法Random Forest 或者 U-Net or analyze出错
            data:
              type: string
              example: error
        description: 失败
    """

    def base64(imgs):
        res = []
        for idx in range(imgs.shape[2]):
            res.append(img_to_base64(imgs[:, :, idx]))
        return res, str(imgs.shape[2])

    try:
        response_object = {}
        json = request.get_json()
        dwi_file, adc_file = [bytes, str]
        modelType = json['backmodel']
        dwi_name = json['dwi_file']
        adc_name = json['adc_file']
        id = json['patientID']
        if dwi_name == "" and adc_name == "":
            return failReturn("", "analyze: 输入参数缺失")
        if dwi_name:
            dwi_file = os.path.join(app.config['UPLOAD_FOLDER'], dwi_name)
        if adc_name:
            adc_file = os.path.join(app.config['UPLOAD_FOLDER'], adc_name)

        imgs = load_imgs(adc_file, dwi_file)
        dwi_arr = imgs.get('dwi')
        adc_arr = imgs.get('adc')
        affine = imgs.get('affine')
        if modelType == "Random Forest":
            perf_preds, nonperf_preds, info, size = stage2(perf_model, nonperf_model, perf_clf, nonperf_clf, dwi_arr,
                                                           adc_arr, socketio)
        elif modelType == "U-Net":
            perf_preds, nonperf_preds, info, size = stage1_2(perf_model, nonperf_model, perf_clf, nonperf_clf, dwi_arr,
                                                       adc_arr, socketio)
        else:
            return failReturn("", "analyze出错, 请选择正确算法Random Forest 或者 U-Net")
        perf_res = to_nii(perf_preds, affine)
        save_name1 = "perf_" + uuid.uuid4().hex + ".nii"
        save_path1 = os.path.join(app.config['RESULT_FOLDER'], save_name1)
        perf_res.to_filename(save_path1)
        nonperf_res = to_nii(nonperf_preds, affine)
        save_name2 = "nonperf_" + uuid.uuid4().hex + ".nii"
        save_path2 = os.path.join(app.config['RESULT_FOLDER'], save_name2)
        nonperf_res.to_filename(save_path2)
        doctor = session["user_id"]
        res_to_db = Result(save_name1, save_name2, modelType, id, doctor, dwi_name, adc_name, float(info), float(size))
        db.session.add(res_to_db)
        db.session.commit()
        perf_preds[perf_preds >= 0.2] = 1
        perf_preds[perf_preds < 0.2] = 0
        nonperf_preds[nonperf_preds >= 0.2] = 1
        nonperf_preds[nonperf_preds < 0.2] = 0
        response_object['perf_res_imgs'], response_object['perf_res_slices'] = base64(perf_preds)
        response_object['nonperf_res_imgs'], response_object['nonperf_res_slices'] = base64(nonperf_preds)
        response_object['res_path1'] = save_name1
        response_object['res_path2'] = save_name2
        response_object['info'] = round(info, 2)
        response_object['size'] = round(size, 2)
        emailSent(
            "分析结果已得出，详情请查看网关。\n patientId: " + str(id) + "res_path1" + save_name1 + "res_path2" +
            save_name2 + "info: " + str(info), "Analyze Result")
        return successReturn(response_object, "analyze: 分析成功")
    except Exception as e:
        return failReturn(format(e), "analyze出错")


@app.route("/api/download/uploadFile/<path:filename>", methods=['GET'])
@login_required
@cross_origin()
def download_file1(filename):
    """
    下载目录一
    :param filename:

    ---
    tags:
      - model_svr API
    parameters:
      - name: filename
        in: path
        type: string
        description: filename
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      fail:
        description: 下载失败
      success:
        description: 下载成功
    """
    try:
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename)
    except Exception as e:
        return failReturn(format(e), "download1出错")


@app.route("/api/download/resultFile/<path:filename>", methods=['GET'])
@login_required
@cross_origin()
def download_file2(filename):
    """
    下载目录二
    :param filename:

    ---
    tags:
      - model_svr API
    parameters:
      - name: filename
        in: path
        type: string
        description: filename
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      fail:
        description: 下载失败
      success:
        description: 下载成功
    """
    try:
        return send_from_directory(app.config['RESULT_FOLDER'], filename)
    except Exception as e:
        return failReturn(format(e), "download1出错")


# def _get_fix_list():
#     """
#     获取结果信息列表
#       res
#     """
#     doctor = _get_current_user()
#     if doctor.userType == 1:
#         results = Result.query.all()
#         res = []
#         for r in results:
#             if r.realimg and r.roi:
#                 users = User.query.filter_by(id=r.doctor_id).first()
#                 patients = Patient.query.filter_by(id=r.patient_id).first()
#                 res.append({"id": r.id, "modelType": r.modeltype,
#                             "doctorName": users.realname,
#                             "patientName": patients.username
#                             })
#         return res
#
#     else:
#         return "not allowed"
#
#
# @app.route('/api/getResultList', methods=['GET'])
# @login_required
# @cross_origin()
# def get_result_list():
#     """
#     获取结果信息列表
#       json
#     ---
#     tags:
#       - model_svr API
#     parameters:
#       - name: Authorization
#         in: header
#         type: string
#         required: true
#         description: token
#     responses:
#       success:
#         schema:
#           type: object
#           properties:
#             status:
#               type: string
#               example: success.
#             msg:
#               type: string
#               example: 获取成功
#             data:
#               type: object
#               properties:
#                 res:
#                     type: array
#                     items:
#                         type: object
#                         properties:
#                             id:
#                                 type: integer
#                                 example: 1
#                             modelType:
#                                 type: string
#                                 example: modelType
#                             doctorName:
#                                 type: string
#                                 example: doctorName
#                             patientName:
#                                 type: string
#                                 example: patientName
#         description: 成功
#       fail:
#         schema:
#           type: object
#           properties:
#             status:
#               type: string
#               example: fail.
#             msg:
#               type: string
#               example: getFixList出错 or 权限不足
#             data:
#               type: string
#               example: error or not allowed
#         description: 失败
#     """
#     try:
#         res = _get_fix_list()
#         if res == "not allowed":
#             return failReturn("not allowed", "getFixList: 权限不足")
#         elif res:
#             return successReturn({'res': res}, "getFixList: 获取成功")
#         else:
#             return failReturn("", "getFixList: 获取失败")
#     except Exception as e:
#         return failReturn(format(e), "getFixList出错")


@app.route('/api/ROI', methods=['POST'])
@login_required
@cross_origin()
def ROI_upload():
    """
    脑部梗死区上传
      json
    ---
    tags:
      - model_svr API
    parameters:
      - name: file
        in: formData
        type: file
        description: file
      - name: resultID
        in: formData
        type: integer
        description: resultID
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      success:
        schema:
          type: object
          properties:
            status:
              type: string
              example: success.
            msg:
              type: string
              example: 上传成功
            data:
              type: object
              properties:
                roi:
                  type: object
                  properties:
                    roi_file:
                        type: string
                        example: roi_file
                    roi_imgs:
                        type: array
                        items:
                            type: string
                            example: [data:image/jpg;base64 xxxxxx, ···]
                    roi_slices:
                        type: string
                        example: 21
        description: 成功
      fail:
        schema:
          type: object
          properties:
            status:
              type: string
              example: fail.
            msg:
              type: string
              example: ROI出错 or 上传失败
            data:
              type: string
              example: error
        description: 失败
    """
    try:
        response_object = {}
        file = request.files['file']
        uploadname = secure_filename(file.filename)
        resultID = request.form['resultID']
        filename = uuid.uuid4().hex + '_' + uploadname
        save_path = app.config['UPLOAD_FOLDER']
        os.makedirs(save_path, exist_ok=True)
        save_file = os.path.join(save_path, filename)
        r = Result.query.filter_by(id=resultID).first()
        if not r:
            return failReturn("", "ROI: 上传失败")
        else:
            if r.roi:
                os.remove(os.path.join(save_path, r.roi))
            r.roi = filename
            db.session.commit()
            file.save(save_file)
        roi = {'roi_file': save_file, 'roi_imgs': get_all_slice(save_file)[0],
               'roi_slices': get_all_slice(save_file)[1]}
        response_object['roi'] = roi
        return successReturn(response_object, "ROI: 上传成功")
    except Exception as e:
        return failReturn(format(e), "ROI出错")


@app.route('/api/realimg', methods=['POST'])
@login_required
@cross_origin()
def realimg_upload():
    """
    真实图像上传
      json
    ---
    tags:
      - model_svr API
    parameters:
      - name: file
        in: formData
        type: file
        description: file
      - name: resultID
        in: formData
        type: integer
        description: resultID
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      success:
        schema:
          type: object
          properties:
            status:
              type: string
              example: success.
            msg:
              type: string
              example: 上传成功
            data:
              type: object
              properties:
                realimg:
                  type: object
                  properties:
                    realimg_file:
                        type: string
                        example: realimg_file
                    realimg_imgs:
                        type: array
                        items:
                            type: string
                            example: [data:image/jpg;base64 xxxxxx, ···]
                    realimg_slices:
                        type: string
                        example: 21
        description: 成功
      fail:
        schema:
          type: object
          properties:
            status:
              type: string
              example: fail.
            msg:
              type: string
              example: realimg出错 or 上传失败
            data:
              type: string
              example: error
        description: 失败
    """
    try:
        response_object = {}
        file = request.files['file']
        uploadname = secure_filename(file.filename)
        resultID = request.form['resultID']
        filename = uuid.uuid4().hex + '_' + uploadname
        save_path = app.config['UPLOAD_FOLDER']
        os.makedirs(save_path, exist_ok=True)
        save_file = os.path.join(save_path, filename)
        r = Result.query.filter_by(id=resultID).first()
        if not r:
            return failReturn("", "realimg: 上传失败")
        else:
            if r.realimg:
                os.remove(os.path.join(save_path, r.realimg))
            r.realimg = filename
            db.session.commit()
            file.save(save_file)
        realimg = {'realimg_file': save_file, 'realimg_imgs': get_all_slice(save_file)[0],
                   'realimg_slices': get_all_slice(save_file)[1]}
        response_object['realimg'] = realimg
        return successReturn(response_object, "realimg: 上传成功")
    except Exception as e:
        return failReturn(format(e), "realimg出错")


def _get_roi_and_real(id):
    result = Result.query.filter_by(id=id).first()
    if not result:
        return None, None
    save_path = app.config['UPLOAD_FOLDER']
    realimg = os.path.join(save_path, result.realimg)
    roi = os.path.join(save_path, result.roi)
    return realimg, roi


@app.route('/api/getRoiAndReal', methods=['POST'])
@login_required
@cross_origin()
def get_roi_and_real():
    """
      获取roi和realImg
    ---
    tags:
      - model_svr API
    parameters:
      - name: body
        in: body
        required: true
        schema:
          id: 获取并修正
          required:
            - resultID
          properties:
            resultID:
              type: integer
              description: resultID
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      success:
        schema:
          type: object
          properties:
            status:
              type: string
              example: success.
            msg:
              type: string
              example: 获取成功
            data:
              type: object
              properties:
                realimg:
                    type: object
                    properties:
                        realimg_file:
                            type: string
                            example: realimg_file
                        realimg_imgs:
                            type: array
                            items:
                                type: string
                                example: [data:image/jpg;base64 xxxxxx, ···]
                        realimg_slices:
                            type: string
                            example: 21
                roi:
                    type: object
                    properties:
                        roi_file:
                            type: string
                            example: roi_file
                        roi_imgs:
                            type: array
                            items:
                                type: string
                                example: [data:image/jpg;base64 xxxxxx, ···]
                        roi_slices:
                            type: string
                            example: 21
        description: 成功
      fail:
        schema:
          type: object
          properties:
            status:
              type: string
              example: fail.
            msg:
              type: string
              example: getInpFix出错 or 获取失败
            data:
              type: string
              example: error
        description: 失败
    """
    try:
        response_object = {}
        json = request.get_json()
        resultID = json['resultID']
        realimg, roi = _get_roi_and_real(resultID)
        if realimg or roi:
            if realimg:
                realimg = {'realimg_file': realimg, 'realimg_imgs': get_all_slice(realimg)[0],
                           'realimg_slices': get_all_slice(realimg)[1]}
                response_object['realimg'] = realimg
            if roi:
                roi = {'roi_file': roi, 'roi_imgs': get_all_slice(roi)[0],
                       'roi_slices': get_all_slice(roi)[1]}
                response_object['roi'] = roi
        else:
            return failReturn("", "getInpFix: 获取失败")
        return successReturn(response_object, "getInpFix: 获取成功")
    except Exception as e:
        return failReturn(format(e), "getInpFix出错")


def _eval(gt, pred, dwi):
    """
    评测
    :param gt:
    :param pred:
    :param dwi:
      accuracy, specifity, sensitivity, auc
    """
    idx = np.where(dwi.flatten() > 1000)
    preds = pred.flatten()[idx]
    gts = gt.flatten()[idx]
    tn, fp, fn, tp = metrics.confusion_matrix(gts, preds > 0.15).ravel()
    sensitivity = tp / (tp + fn)
    specifity = tn / (fp + tn)
    auc = metrics.roc_auc_score(gts, preds)
    accuracy = metrics.accuracy_score(gts, preds > 0.15)
    return accuracy, specifity, sensitivity, auc


@app.route('/api/eval', methods=['POST'])
@login_required
@cross_origin()
def eval():
    """
    发起评测获得结果
      json
    ---
    tags:
      - model_svr API
    parameters:
      - name: body
        in: body
        required: true
        schema:
          id: 发起评测获得结果
          required:
            - resultID
          properties:
            resultID:
              type: integer
              description: resultID
            dataset:
              type: integer
              description: 0为perfusion数据模型，其他为Non-perfusion数据模型
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      success:
        schema:
          type: object
          properties:
            status:
              type: string
              example: success.
            msg:
              type: string
              example: 发起评测成功
            data:
              type: object
              properties:
                eval:
                    type: string
                    example: 与真实结果相比，perfusion数据模型预测结果准确率为1，特异度为1，灵敏度为1，AUC为1
        description: 成功
      fail:
        schema:
          type: object
          properties:
            status:
              type: string
              example: fail.
            msg:
              type: string
              example: 发起评测失败 or roi参数缺失 or dwi参数缺失 or perf参数缺失 or nonperf参数缺失 or eval出错
            data:
              type: string
              example: error or not allowed
        description: 失败
    """
    try:
        response_object = {}
        resultID = request.get_json()['resultID']
        dataset = request.get_json()['dataset']
        res = Result.query.filter_by(id=resultID).first()
        if not res:
            return failReturn("", "eval: 发起评测失败")
        roi = res.roi
        if roi is None:
            return failReturn("", "eval: roi参数缺失")
        roi = nib.load(os.path.join(app.config['UPLOAD_FOLDER'], roi)).get_fdata()
        roi = np.squeeze(roi)
        dwi = res.dwi_name
        if dwi is None:
            return failReturn("", "eval: dwi参数缺失")
        dwi = nib.load(os.path.join(app.config['UPLOAD_FOLDER'], dwi)).get_fdata()
        dwi = np.squeeze(dwi)
        perf = res.filename1
        if perf is None:
            return failReturn("", "eval: perf参数缺失")
        perf = nib.load(os.path.join(app.config['RESULT_FOLDER'], perf)).get_fdata()
        perf = np.squeeze(perf)
        nonperf = res.filename2
        if nonperf is None:
            return failReturn("", "eval: nonperf参数缺失")
        nonperf = nib.load(os.path.join(app.config['RESULT_FOLDER'], nonperf)).get_fdata()
        nonperf = np.squeeze(nonperf)
        if dataset == 0:
            accuracy, specifity, sensitivity, auc = _eval(roi, perf, dwi)
            response_object['eval'] = '与真实结果相比，perfusion数据模型预测结果准确率为{}，特异度为{}，灵敏度为{}，AUC为{}'. \
                format(round(accuracy, 2), round(specifity, 2), round(sensitivity, 2), round(auc, 2))
        else:
            accuracy, specifity, sensitivity, auc = _eval(roi, nonperf, dwi)
            response_object['eval'] = '与真实结果相比，Non-perfusion数据模型预测结果准确率为{}，特异度为{}，灵敏度为{}，AUC为{}'. \
                format(round(accuracy, 2), round(specifity, 2), round(sensitivity, 2), round(auc, 2))
        return successReturn(response_object, "eval: 发起评测成功")
    except Exception as e:
        return failReturn(format(e), "eval出错")


def _get_report(doctor, patient, result):
    document = Document()
    title1 = document.add_heading('浙江大学第一附属医院', 0)
    title1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title2 = document.add_heading('脑卒中诊疗辅助报告', 0)
    title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    d = document.add_paragraph()
    run1 = d.add_run('责任医生：')
    run1.font.size = Pt(23)
    run2 = d.add_run(doctor.username)
    run2.font.size = Pt(23)
    run2.font.color.rgb = RGBColor(178, 34, 34)
    document.add_paragraph().add_run('报告生成时间：' + str(datetime.now().strftime('%y-%m-%d %H:%M:%S'))).font.size = Pt(18)
    document.add_heading('病人基本信息', level=1)
    if patient.sex == 1:
        sex = "男"
    else:
        sex = "女"
    document.add_paragraph(
        "病人：" + patient.username + "，年龄：" + str(patient.age) + "，性别：" + sex + "，病例本编号：" + str(patient.record_id) + "。",
        style='List Bullet')
    document.add_paragraph("病人病情：" + str(patient.info) + "。", style='List Bullet')
    document.add_paragraph("诊疗结果：" + str(patient.result) + "。", style='List Bullet')
    document.add_paragraph("脑卒中类别判断：" + str(patient.cva) + "。当前状态：" + str(patient.state) + "。", style='List Bullet')
    document.add_paragraph("初次看病时间：" + str(patient.create_time) + "。最近一次看病时间：" + str(patient.update_time) + "。",
                           style='List Bullet')
    _img_process(result.adc_name)
    _img_process(result.dwi_name)
    _img_process(result.filename1)
    _img_process(result.filename2)
    document.add_page_break()
    document.add_paragraph("ADC图像", style='List Bullet')
    document.add_picture(os.path.join(app.config['DOC_FOLDER'], 'pic', 'adc.png'))
    document.add_page_break()
    document.add_paragraph("DWI图像", style='List Bullet')
    document.add_picture(os.path.join(app.config['DOC_FOLDER'], 'pic', 'dwi.png'))
    document.add_page_break()
    document.add_paragraph("perf图像", style='List Bullet')
    document.add_picture(os.path.join(app.config['DOC_FOLDER'], 'pic', 'perf.png'))
    document.add_page_break()
    document.add_paragraph("nonperf图像", style='List Bullet')
    document.add_picture(os.path.join(app.config['DOC_FOLDER'], 'pic', 'nonperf.png'))
    filename = str(uuid.uuid4()) + '_res.docx'
    document.save(os.path.join(app.config['DOC_FOLDER'], filename))
    return filename


def _img_process(imgName):
    name = imgName.split("_")
    if name[0].lower() == "adc" or name[0].lower() == "dwi":
        path = os.path.join(app.config['UPLOAD_FOLDER'], imgName)
        file = nib.load(path)
        img_arr = file.dataobj[:, :, 10]
        plt.imshow(img_arr, cmap='gray')
        if name[0].lower() == "adc":
            adc_path = os.path.join(app.config['DOC_FOLDER'], 'pic', 'adc.png')
            plt.savefig(adc_path)
            plt.close(adc_path)
        elif name[0].lower() == "dwi":
            dwi_path = os.path.join(app.config['DOC_FOLDER'], 'pic', 'dwi.png')
            plt.savefig(dwi_path)
            plt.close(dwi_path)
    elif name[0].lower() == "perf" or name[0].lower() == "nonperf":
        path = os.path.join(app.config['RESULT_FOLDER'], imgName)
        file = nib.load(path)
        img_arr = file.dataobj[:, :, 10]
        plt.imshow(img_arr, cmap='gray')
        if name[0].lower() == "perf":
            perf_path = os.path.join(app.config['DOC_FOLDER'], 'pic', 'perf.png')
            plt.savefig(perf_path)
            plt.close(perf_path)
        elif name[0].lower() == "nonperf":
            nonperf_path = os.path.join(app.config['DOC_FOLDER'], 'pic', 'nonperf.png')
            plt.savefig(nonperf_path)
            plt.close(nonperf_path)


@app.route('/api/getReport', methods=['POST'])
@login_required
@cross_origin()
def get_report():
    """
    获取analyze报告
      json
    ---
    tags:
      - model_svr API
    parameters:
      - name: body
        in: body
        required: true
        schema:
          id: 获取analyze报告
          required:
            - resultID
          properties:
            resultID:
              type: integer
              description: resultID
      - name: Authorization
        in: header
        type: string
        required: true
        description: token
    responses:
      success:
        schema:
          type: file
          example: file
        description: 成功
      fail:
        schema:
          type: object
          properties:
            status:
              type: string
              example: fail.
            msg:
              type: string
              example: getReport出错 or 权限不足无法查看
            data:
              type: string
              example: error or not allowed
        description: 失败
    """
    try:
        resultID = request.get_json()['resultID']
        result = Result.query.filter_by(id=resultID).first()
        user = _get_current_user()
        if user.userType != 1 and user.id != result.doctor_id:
            return failReturn("", "getReport: 权限不足无法查看")
        patient = Patient.query.filter_by(id=result.patient_id).first()
        doctor = User.query.filter_by(id=result.doctor_id).first()
        filename = _get_report(doctor, patient, result)
        response = make_response(send_from_directory(app.config['DOC_FOLDER'], filename, as_attachment=True))
        return response
    except Exception as e:
        return failReturn(format(e), "getReport出错")


if __name__ == '__main__':
    def after_request(resp):
        resp.headers['Access-Control-Allow-Credentials'] = 'true'
        return resp


    perf_model, nonperf_model = stage1_init()
    perf_clf, nonperf_clf = stage2_init()
    app.after_request(after_request)
    socketio.run(app, host='0.0.0.0', port=5051, debug=True)
